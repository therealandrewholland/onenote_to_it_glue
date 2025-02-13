import os
import re
import time
import uuid
import json
import string
import shutil
import zipfile
import traceback
from urllib.parse import quote
import win32com.client
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE


def getURLOrPath(client, client_prospects_path):
    global manual_check_needed
    global client_name
    
    url_or_path = None
    file_name = ""

    if client == "HH Marshall, LLC":
        client_folder = "HH Cassopolis LLC"
    else:
        client_folder = client
    onenote_path, mid_string = os.path.join(client_prospects_path, f"{client_folder}\\Operations\\Documentation\\OneNote"), "Operations/Documentation/OneNote"
    if not os.path.exists(onenote_path):
        onenote_path, mid_string = os.path.join(client_prospects_path, f"{client_folder}\\Documentation\\OneNote"), "Documentation/OneNote"
    if not os.path.exists(onenote_path):
        print(f"No documentation folders found in {client_prospects_path}\\{client_folder}")
        return url_or_path, file_name
    
    file_list = os.listdir(onenote_path)

    # Get .url or .one. Does basic check to see if the start of the client and file names are the same
    notebook_file_list = []
    for file in file_list:
        if (file.endswith('.url') or file.endswith('.one')) and (client[:5].lower() == file[:5].lower()):
            notebook_file_list.append(file)
    
    if len(notebook_file_list) > 1:
        manual_check_needed[client_name].append(f"Many valid .url or .one files found in {onenote_path}: {notebook_file_list}")

        file_list_with_paths = []
        for file in file_list:
            file_list_with_paths.append(os.path.join(onenote_path, file))

        #get newest file if there are mutliple    
        full_path = max(file_list_with_paths, key=os.path.getmtime)
        file_name = os.path.basename(full_path)
        
    elif len(notebook_file_list) == 0:
        manual_check_needed[client_name].append(f"No valid .url or .one files found in {onenote_path}: {file_list}")
    else:
        file_name = notebook_file_list[0]

    if file_name.endswith('.url'):
        #Remove ".url" so it will work in full URL later
        file_name = file_name[:-len('.url')]

        encoded_file_name = quote(file_name)
        encoded_client_folder = quote(client_folder)

        url_or_path = f"onenote:https://aunalytics.sharepoint.com/sites/Aunalytics2/Document%20Library/Clients%20&%20Prospects/{encoded_client_folder}/{mid_string}/{encoded_file_name}/"
    elif file_name.endswith('.one'):
        url_or_path = os.path.join(onenote_path, file_name)
 
    return url_or_path, file_name

def getNotebookID(path_or_url, url_file=False, notebook_name=None):
        app = win32com.client.dynamic.Dispatch('OneNote.Application')

         #wait to finish opening notebook before continuing
        if url_file:
            #open url in OneNote
            app.NavigateToUrl(path_or_url, False)

            #wait to finish opening notebook before continuing
            time.sleep(10)
            while True:
                time.sleep(3)

                try:
                    open_notebooks = app.GetHierarchy("" , 1)
                    break
                except:
                    continue

            root = ET.fromstring(open_notebooks)
            namespaces = {'one': 'http://schemas.microsoft.com/office/onenote/2013/onenote'}
            all_notebooks = root.findall('one:Notebook', namespaces)
            for notebook in all_notebooks:
                if notebook.get('name') == notebook_name:
                    notebook_id = notebook.get('ID')
                    break
                
        else:      
            app = win32com.client.dynamic.Dispatch('OneNote.Application')

            #store older open notebooks / sections
            old_open_notebooks = app.GetHierarchy("" , 1)
            
            app.OpenHierarchy(path_or_url, "", "", 0)

            #store newer open notebooks / sections after opening another
            new_open_notebooks = app.GetHierarchy("" , 1)

            #get ID from the difference between old / new (this lets it handle both notebooks and sections)
            result = ""
            for i in range(len(new_open_notebooks)):
                if i >= len(old_open_notebooks) or new_open_notebooks[i] != old_open_notebooks[i]:
                    result += new_open_notebooks[i]

            # Remove uneeded information around id
            notebook_id = result.split("\"")[1]

        return notebook_id

#url can be path to .one or formatted URL that starts with "onenote:"
def getID(url_or_path, notebook_name):
    global manual_check_needed
    global client_name
    
    app = win32com.client.dynamic.Dispatch('OneNote.Application')
    
    if url_or_path.startswith("onenote:"):
        #This will fetch the highest ID for the .one file
        notebook_id = getNotebookID(url_or_path, url_file=True, notebook_name=notebook_name)
    else:
        #for akzo. This will fetch the highest ID for the .one file
        notebook_id = getNotebookID(url_or_path, url_file=False)

    #wait for all sections to sync
    time.sleep(10)
    main_sections = app.GetHierarchy(notebook_id, 1)
    root = ET.fromstring(main_sections)
    namespaces = {'one': 'http://schemas.microsoft.com/office/onenote/2013/onenote'}
    all_sections = root.findall('one:Section', namespaces)

    # This dictionary will store the "Production" section id and only overwrite it if there is a "Production 2" or some variation
    id_list = {
        "name": None,
        "ID": None
    }
    
    production_list = []
    
    for section in all_sections:
        name = section.get('name')
        #for akzo. If these is only one section, grabs its id instead of the "Production" section
        if "Production" in name or len(all_sections) == 1:
            production_list.append(name)
            if id_list["name"] == None or "2" in name:
                id_list["name"] = name
                id_list["ID"] = section.get('ID')

    if len(production_list) > 1:
        manual_check_needed[client_name].append(f"Many \"Production\" tabs found in client notebook: {production_list}")
                
    target_section_id = id_list["ID"]
    
    return target_section_id, notebook_id

def createDoc(target_section_id, client):
    app = win32com.client.dynamic.Dispatch('OneNote.Application')
    docx_path = os.path.join(os.getcwd(), f"temp\\{client}.docx")
    app.Publish(target_section_id, docx_path, 5, "")


def getHeadingsFromNotebook(target_section_id):
    app = win32com.client.dynamic.Dispatch('OneNote.Application')

    # Get all notebook info
    notebooks_xml = app.GetHierarchy(target_section_id, 4)
    root = ET.fromstring(notebooks_xml)
    namespace = {'one': 'http://schemas.microsoft.com/office/onenote/2013/onenote'}

    #Used for checking full notebook as raw .xml
    #print(notebooks_xml)

    #Track heading text as key and "Heading" level as value
    heading_levels_dict = {}
    for page in root.findall('one:Page', namespace):
        page_name = page.get('name')
        page_level = page.get('pageLevel')

        heading_levels_dict[page_name] = page_level

    return heading_levels_dict


def getHeadingsFromDocx(document, target_section_id):
    heading_tracking_listA = []
    heading_tracking_listB =[]
    heading_tracking_number = 0
    heading_levels_dict = getHeadingsFromNotebook(target_section_id)     
        
    # Iterate through .docx to identify sections / subsetions by 
    for paragraph in document.paragraphs:
        if paragraph.style.name in ['Heading 1', 'Heading 2', 'Heading 3']:
            paragraph.style = "Normal"
        text = paragraph.text

        if (text in heading_levels_dict):
            heading_tracking_listA.append(heading_tracking_number)
        if (':' in text and (text.endswith(" AM") or text.endswith(" PM"))):
            heading_tracking_listB.append(heading_tracking_number-2)
            
        heading_tracking_number += 1

        #Used for checking value of every parahraph in .docx
        #print(paragraph.text)
        #print("- - - - - - - - - - -")
    heading_tracking_list = list(set(heading_tracking_listA) & set(heading_tracking_listB))
    
    return heading_tracking_list, heading_levels_dict


def formatHeadings(client, target_section_id):
    global manual_check_needed
    global client_name
    
    document = Document(f"temp\\{client}.docx")

    if len(document.paragraphs) == 1:
        failed_list = findCorruptSections(target_section_id)
        manual_check_needed[client_name].append(f"Error exporting from OneNote file. Recreate corrupted sections: {failed_list}")

    heading_tracking_list, heading_levels_dict = getHeadingsFromDocx(document, target_section_id)

    # For each paragraph that needs to be a heading, add level (i.e. "Heading 1", "Heading 2")
    for i in heading_tracking_list:
        paragraph = document.paragraphs[i]
        key = paragraph.text
        heading_level = heading_levels_dict[key]
        heading_convert = {
            '1': 'Heading 1',
            '2': 'Heading 2',
            '3': 'Heading 3'
        }

        #Used for checking headings and their level
        #print(key)
        #print(heading_convert[heading_level])
        #print("- -- -- -- -- -- -")

        document.paragraphs[i].style = heading_convert[heading_level]

    # .docx files show weird when uploaded to IT Glue
    document.save(f"temp\\final_docx\\{client}.docx")
    os.remove(f"temp\\{client}.docx")

# Used for checking existing .docx style info
"""
def getStyleInfo():
    doc = Document('test.docx')

    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            style_name = style.name
            print(style_name)

            paragraph_format = style.paragraph_format
            
            for attribute in dir(paragraph_format):
                try:
                    if not attribute.startswith('__'):
                        value = getattr(paragraph_format, attribute)
                        print(f"{attribute}: {value}")
                except:
                    pass

            print("- - - - - - - - -")
"""

# TO DO
def extractDocx(document_name, extract_folder_name):
    with zipfile.ZipFile(document_name, 'r') as zip_ref:
        zip_ref.extractall(extract_folder_name)

    return os.path.join(extract_folder_name, 'word', 'styles.xml')

def rezipDocx(document_name, extract_folder_name):
    with zipfile.ZipFile(document_name, 'w', zipfile.ZIP_DEFLATED) as docx_zip:
        for foldername, subfolders, filenames in os.walk(extract_folder_name):
            for filename in filenames:
                filepath = os.path.join(foldername, filename)
                arcname = os.path.relpath(filepath, extract_folder_name)
                docx_zip.write(filepath, arcname)   

#pull Headings styles from template .Zip and add them to original
def createStyles(client):
    document = Document(f"temp\\{client}.docx")
    
    template_document_xml_path = extractDocx("template.docx", "temp\\extracted_contents\\template")
    document_xml_path = extractDocx(f"temp\\{client}.docx", "temp\\extracted_contents\\test")

    os.remove(document_xml_path)
    os.remove(f"temp\\{client}.docx")
    shutil.copy2(template_document_xml_path, document_xml_path)

    rezipDocx(f"temp\\{client}.docx", "temp\\extracted_contents\\test")
    
    document = Document(f"temp\\{client}.docx")

    cleanup_list = ["temp\\extracted_contents\\template", "temp\\extracted_contents\\test"]
    for directory in cleanup_list:
        shutil.rmtree(directory)

# Not really needed. Used for checking and formatting client .docx files before uploading
def format_client(client):
    client_2 = client.replace(',', '').lower().replace('inc', '').replace('.', '').replace(',', '').replace(' ', '')
    return client_2

def crossCheckLists(all_client_list, pod_list):
    formatted_list_1 = []
    formatted_list_2 = []
    client_dict = {}

    pod_list += ["City of Hastings", "Lineage Logistics", "Payroc WorldAccess LLC"]

    for client in all_client_list:
        formatted_client = format_client(client)
        formatted_list_1.append(formatted_client)
        client_dict[formatted_client] = client

    for client in pod_list:
        formatted_client = format_client(client)
        formatted_list_2.append(formatted_client)


    keys = list(set(formatted_list_1) & set(formatted_list_2))

    return keys, client_dict


"""
def getBodeList():
    all_client_list = os.listdir("formatted_docs_2\\")
    new_all_client_list = []
    for client in all_client_list:
        if not client.startswith('~$'):
            new_all_client_list.append(client[:-5])
        
    new_all_client_list.append("ABI Attachments")
    new_all_client_list.append("Adventist Frontier Missions")

    return new_all_client_list
"""

#Fixes issue with formatted docx not working in IT Glue by opening them in Word and converting to .doc
def convertToDoc(client, formatted_client):
    current_directory = os.getcwd()
    
    word = win32com.client.dynamic.Dispatch("Word.Application")
    word.Visible = False
    
    document = word.Documents.Open(f"{current_directory}\\temp\\final_docx\\{formatted_client}.docx")

    #Adds client name to start of file so the file name shows correctly in IT Glue
    range = document.Content
    range.InsertBefore(client + ".OneNote\n")
    range = document.Range(0, len(client) + 1)
    range.Style = "Heading 1"
    
    document.SaveAs(f"{current_directory}\\formatted_docs\\{client}.doc", FileFormat=0)

    document.Close()
    word.Quit()
    os.remove(f"temp\\final_docx\\{formatted_client}.docx")


def findCorruptSections(target_section_id):
    app = win32com.client.dynamic.Dispatch('OneNote.Application')
    notebook_info = app.GetHierarchy(target_section_id, 4)
    root = ET.fromstring(notebook_info)
    namespace = "{http://schemas.microsoft.com/office/onenote/2013/onenote}Page"

    os.makedirs("temp\\find_corrupt_section\\")
    failed_list = []
    for page in root.iter(namespace):
        page_name = page.get('name')
        formatted_page_name = page_name.lower().translate(str.maketrans('', '', string.punctuation))
        page_id = page.get('ID')
        docx_path = os.path.join(os.getcwd(), f"temp\\find_corrupt_section\\{formatted_page_name}.one")

        try:
            app.Publish(page_id, docx_path, 0, "")
        except Exception as e:
            failed_list.append(page_name)

    shutil.rmtree("temp\\find_corrupt_section\\")
            
    return failed_list
    
def Main():
    global manual_check_needed
    global client_name
    
    debugging = False
    client_prospects_path = os.path.join(os.environ['OneDrive'], "Clients & Prospects")
    all_client_list = os.listdir(client_prospects_path)
    notebook_id = None

    # change the client and debugging boolean to enable debugging for specific clients 
    if debugging:
        all_client_list = ["HH Marshall, LLC", "Primary Care Partners of South Bend, LLC", "StresCore, Inc", "Transpo"]
        
    #Can be used to import a list from .txt file
    
    with open("Client_List.txt", 'r') as file:
        data = file.read()
        client_list = eval(data)

    keys, client_dictionary = crossCheckLists(all_client_list, client_list)

    for key in keys:
        client = client_dictionary[key]
        try:
            formatted_client = client.replace(',', '').lower().replace(' ', '_')
            client_name = client
            manual_check_needed[client_name] = []

            if not os.path.exists(f"formatted_docs\\{client}.doc"):
                url_or_path, file_name = getURLOrPath(client, client_prospects_path)

                if not url_or_path == None:
                    target_section_id, notebook_id = getID(url_or_path, file_name)
                    createDoc(target_section_id, formatted_client)
                    createStyles(formatted_client)
                    formatHeadings(formatted_client, target_section_id)
                    if len(manual_check_needed[client_name]) == 0:
                        convertToDoc(client, formatted_client)

        #Catch and print all errors for failed clients
        except Exception as e:
            print(f"{client} failed. Skipping. Error message:")
            print(e)
            traceback.print_exc()
            continue
        
        finally:
            #close any opened notebooks, ideally even if the client fails
            if not notebook_id == None:
                app = win32com.client.dynamic.Dispatch('OneNote.Application')
                try:
                    app.CloseNotebook(notebook_id)
                except:
                    print(f"Unable to close notebook or section for {client}")
                    

        if manual_check_needed:
            sorted_dict = dict(sorted(manual_check_needed.items()))
            with open("need_review\\manual_check_needed.json", 'w') as json_file:
                json.dump(sorted_dict, json_file)

#Global variable used to store what clients need review
manual_check_needed = {}
client_name = ""

Main()


